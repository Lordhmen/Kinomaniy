import sys
import os
import docx
import requests

from PySide2 import QtGui
from PySide2.QtGui import QIntValidator
from docx.shared import Pt
from PySide2.QtCore import (Qt, QPoint)
from PySide2.QtWidgets import *

from Functions.ui_functions import *
from Films.script import *
import AddZapis.add_zapis
import User.script_user


class AddZApis(AddZapis.add_zapis.ADDZAPIS):
    def __init__(self):
        super().__init__()
        # self.setupUi(self)


class User(User.script_user.USER):
    def __init__(self):
        super().__init__()
        # self.setupUi(self)


class MainWindows(QMainWindow):
    def __init__(self, parent=None):
        global T1, root
        QMainWindow.__init__(self, parent)
        self.addzapis = None
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.dragPos = QPoint()
        self.ROOT()
        self.usser = User()
        self.FILMS()
        self.films_mesta_obr()
        self.gurnal_start()
        self.inf()
        self.ui.comboBox.activated.connect(self.inf_film)
        self.ui.comboBox_2.activated.connect(self.films_mesta_obr_one)
        self.ui.comboBox_4.activated.connect(self.films_mesta_obr)
        self.ui.pushButton_98.clicked.connect(self.printer)
        self.ui.pushButton.clicked.connect(self.del_prokat)
        self.ui.pushButton_2.clicked.connect(self.add_prokat)
        self.ui.pushButton_94.clicked.connect(self.add_gurnal_zapis)
        self.ui.pushButton_95.clicked.connect(self.del_gurnal_zapis)
        self.ui.pushButton_93.clicked.connect(self.seve_gurnal)
        self.ui.pushButton_96.clicked.connect(self.inf)
        self.ui.pushButton_97.clicked.connect(self.zapross)

        # Обработка кнопок Места в зале
        self.ui.pushButton_3.clicked.connect(lambda: self.obraboka_mest(1, 1, 3))
        self.ui.pushButton_4.clicked.connect(lambda: self.obraboka_mest(1, 2, 4))
        self.ui.pushButton_5.clicked.connect(lambda: self.obraboka_mest(1, 3, 5))
        self.ui.pushButton_6.clicked.connect(lambda: self.obraboka_mest(1, 4, 6))
        self.ui.pushButton_7.clicked.connect(lambda: self.obraboka_mest(1, 5, 7))
        self.ui.pushButton_8.clicked.connect(lambda: self.obraboka_mest(1, 6, 8))
        self.ui.pushButton_9.clicked.connect(lambda: self.obraboka_mest(1, 7, 9))
        self.ui.pushButton_10.clicked.connect(lambda: self.obraboka_mest(1, 8, 10))
        self.ui.pushButton_11.clicked.connect(lambda: self.obraboka_mest(1, 9, 11))
        self.ui.pushButton_12.clicked.connect(lambda: self.obraboka_mest(1, 10, 12))
        self.ui.pushButton_13.clicked.connect(lambda: self.obraboka_mest(1, 15, 13))
        self.ui.pushButton_14.clicked.connect(lambda: self.obraboka_mest(1, 14, 14))
        self.ui.pushButton_15.clicked.connect(lambda: self.obraboka_mest(1, 13, 15))
        self.ui.pushButton_16.clicked.connect(lambda: self.obraboka_mest(1, 12, 16))
        self.ui.pushButton_17.clicked.connect(lambda: self.obraboka_mest(1, 11, 17))
        self.ui.pushButton_18.clicked.connect(lambda: self.obraboka_mest(2, 15, 18))
        self.ui.pushButton_19.clicked.connect(lambda: self.obraboka_mest(2, 4, 19))
        self.ui.pushButton_20.clicked.connect(lambda: self.obraboka_mest(2, 1, 20))
        self.ui.pushButton_21.clicked.connect(lambda: self.obraboka_mest(2, 8, 21))
        self.ui.pushButton_22.clicked.connect(lambda: self.obraboka_mest(2, 9, 22))
        self.ui.pushButton_23.clicked.connect(lambda: self.obraboka_mest(2, 10, 23))
        self.ui.pushButton_24.clicked.connect(lambda: self.obraboka_mest(2, 12, 24))
        self.ui.pushButton_25.clicked.connect(lambda: self.obraboka_mest(2, 13, 25))
        self.ui.pushButton_26.clicked.connect(lambda: self.obraboka_mest(2, 3, 26))
        self.ui.pushButton_27.clicked.connect(lambda: self.obraboka_mest(2, 14, 27))
        self.ui.pushButton_28.clicked.connect(lambda: self.obraboka_mest(2, 2, 28))
        self.ui.pushButton_29.clicked.connect(lambda: self.obraboka_mest(2, 6, 29))
        self.ui.pushButton_30.clicked.connect(lambda: self.obraboka_mest(2, 7, 30))
        self.ui.pushButton_31.clicked.connect(lambda: self.obraboka_mest(2, 11, 31))
        self.ui.pushButton_32.clicked.connect(lambda: self.obraboka_mest(2, 5, 32))
        self.ui.pushButton_33.clicked.connect(lambda: self.obraboka_mest(3, 15, 33))
        self.ui.pushButton_34.clicked.connect(lambda: self.obraboka_mest(3, 4, 34))
        self.ui.pushButton_35.clicked.connect(lambda: self.obraboka_mest(3, 1, 35))
        self.ui.pushButton_36.clicked.connect(lambda: self.obraboka_mest(3, 8, 36))
        self.ui.pushButton_37.clicked.connect(lambda: self.obraboka_mest(3, 9, 37))
        self.ui.pushButton_38.clicked.connect(lambda: self.obraboka_mest(3, 10, 38))
        self.ui.pushButton_39.clicked.connect(lambda: self.obraboka_mest(3, 12, 39))
        self.ui.pushButton_40.clicked.connect(lambda: self.obraboka_mest(3, 13, 40))
        self.ui.pushButton_41.clicked.connect(lambda: self.obraboka_mest(3, 3, 41))
        self.ui.pushButton_42.clicked.connect(lambda: self.obraboka_mest(3, 14, 42))
        self.ui.pushButton_43.clicked.connect(lambda: self.obraboka_mest(3, 2, 43))
        self.ui.pushButton_44.clicked.connect(lambda: self.obraboka_mest(3, 6, 44))
        self.ui.pushButton_45.clicked.connect(lambda: self.obraboka_mest(3, 7, 45))
        self.ui.pushButton_46.clicked.connect(lambda: self.obraboka_mest(3, 11, 46))
        self.ui.pushButton_47.clicked.connect(lambda: self.obraboka_mest(3, 5, 47))
        self.ui.pushButton_48.clicked.connect(lambda: self.obraboka_mest(4, 15, 48))
        self.ui.pushButton_49.clicked.connect(lambda: self.obraboka_mest(4, 4, 49))
        self.ui.pushButton_50.clicked.connect(lambda: self.obraboka_mest(4, 1, 50))
        self.ui.pushButton_51.clicked.connect(lambda: self.obraboka_mest(4, 8, 51))
        self.ui.pushButton_52.clicked.connect(lambda: self.obraboka_mest(4, 9, 52))
        self.ui.pushButton_53.clicked.connect(lambda: self.obraboka_mest(4, 10, 53))
        self.ui.pushButton_54.clicked.connect(lambda: self.obraboka_mest(4, 12, 54))
        self.ui.pushButton_55.clicked.connect(lambda: self.obraboka_mest(4, 13, 55))
        self.ui.pushButton_56.clicked.connect(lambda: self.obraboka_mest(4, 3, 56))
        self.ui.pushButton_57.clicked.connect(lambda: self.obraboka_mest(4, 14, 57))
        self.ui.pushButton_58.clicked.connect(lambda: self.obraboka_mest(4, 2, 58))
        self.ui.pushButton_59.clicked.connect(lambda: self.obraboka_mest(4, 6, 59))
        self.ui.pushButton_60.clicked.connect(lambda: self.obraboka_mest(4, 7, 60))
        self.ui.pushButton_61.clicked.connect(lambda: self.obraboka_mest(4, 11, 61))
        self.ui.pushButton_62.clicked.connect(lambda: self.obraboka_mest(4, 5, 62))
        self.ui.pushButton_63.clicked.connect(lambda: self.obraboka_mest(5, 15, 63))
        self.ui.pushButton_64.clicked.connect(lambda: self.obraboka_mest(5, 4, 64))
        self.ui.pushButton_65.clicked.connect(lambda: self.obraboka_mest(5, 1, 65))
        self.ui.pushButton_66.clicked.connect(lambda: self.obraboka_mest(5, 8, 66))
        self.ui.pushButton_67.clicked.connect(lambda: self.obraboka_mest(5, 9, 67))
        self.ui.pushButton_68.clicked.connect(lambda: self.obraboka_mest(5, 10, 68))
        self.ui.pushButton_69.clicked.connect(lambda: self.obraboka_mest(5, 12, 69))
        self.ui.pushButton_70.clicked.connect(lambda: self.obraboka_mest(5, 13, 70))
        self.ui.pushButton_71.clicked.connect(lambda: self.obraboka_mest(5, 3, 71))
        self.ui.pushButton_72.clicked.connect(lambda: self.obraboka_mest(5, 14, 72))
        self.ui.pushButton_73.clicked.connect(lambda: self.obraboka_mest(5, 2, 73))
        self.ui.pushButton_74.clicked.connect(lambda: self.obraboka_mest(5, 6, 74))
        self.ui.pushButton_75.clicked.connect(lambda: self.obraboka_mest(5, 7, 75))
        self.ui.pushButton_76.clicked.connect(lambda: self.obraboka_mest(5, 11, 76))
        self.ui.pushButton_77.clicked.connect(lambda: self.obraboka_mest(5, 5, 77))
        self.ui.pushButton_78.clicked.connect(lambda: self.obraboka_mest(6, 15, 78))
        self.ui.pushButton_79.clicked.connect(lambda: self.obraboka_mest(6, 4, 79))
        self.ui.pushButton_80.clicked.connect(lambda: self.obraboka_mest(6, 1, 80))
        self.ui.pushButton_81.clicked.connect(lambda: self.obraboka_mest(6, 8, 81))
        self.ui.pushButton_82.clicked.connect(lambda: self.obraboka_mest(6, 9, 82))
        self.ui.pushButton_83.clicked.connect(lambda: self.obraboka_mest(6, 10, 83))
        self.ui.pushButton_84.clicked.connect(lambda: self.obraboka_mest(6, 12, 84))
        self.ui.pushButton_85.clicked.connect(lambda: self.obraboka_mest(6, 13, 85))
        self.ui.pushButton_86.clicked.connect(lambda: self.obraboka_mest(6, 3, 86))
        self.ui.pushButton_87.clicked.connect(lambda: self.obraboka_mest(6, 14, 87))
        self.ui.pushButton_88.clicked.connect(lambda: self.obraboka_mest(6, 2, 88))
        self.ui.pushButton_89.clicked.connect(lambda: self.obraboka_mest(6, 6, 89))
        self.ui.pushButton_90.clicked.connect(lambda: self.obraboka_mest(6, 7, 90))
        self.ui.pushButton_91.clicked.connect(lambda: self.obraboka_mest(6, 11, 91))
        self.ui.pushButton_92.clicked.connect(lambda: self.obraboka_mest(6, 5, 92))

        self.ui.plainTextEdit_2.setPlainText(
            'ПП "Киномания" упрощает работу с базой данных кинотеатра, дает возможность добовлять фильмы и просматривать ранее добавленные. Редактировать состояние фильма. Бронировать места в зале и тд')
        self.ui.plainTextEdit_3.setPlainText(
            'Данный программный продукт «Киномания» разработан курсантам 431 группы Кирсанов Г.В. по специальности: 09.02.03 «Программирование в компьютерных системах»')

        # Установка валидатора
        pintValidator = QIntValidator(self)

        pintLineEdit = self.ui.lineEdit
        pintLineEdit.setValidator(pintValidator)

        self.setWindowTitle('Киномания')
        UIFunctions.labelTitle(self, 'Киномания')
        UIFunctions.labelDescription(self, 'Фильмы')

        startSize = QSize(1089, 729)
        self.resize(startSize)
        self.setMinimumSize(startSize)

        self.ui.btn_toggle_menu.clicked.connect(lambda: UIFunctions.toggleMenu(self, 220, True))

        self.ui.stackedWidget.setMinimumWidth(20)
        UIFunctions.addNewMenu(self, "Фильмы", "btn_report", "url(icons/32x32/film-reel.png)",
                               True)
        UIFunctions.addNewMenu(self, "Места в зале", "btn_bill", "url(icons/32x32/ticket.png)",
                               True)
        UIFunctions.addNewMenu(self, "Киножурнал", "btn_2", "url(icons/32x32/playlist.png)",
                               True)
        UIFunctions.addNewMenu(self, "Запросы", "btn_3", "url(icons/32x32/director.png)",
                               True)
        UIFunctions.addNewMenu(self, "Информация", "btn", "url(icons/32x32/information.png)",
                               False)

        UIFunctions.selectStandardMenu(self, "btn_report")

        self.ui.stackedWidget.setCurrentWidget(self.ui.page_report)
        self.ui.frame_label_top_btns.mouseMoveEvent = self.moveWindow
        UIFunctions.uiDefinitions(self)
        self.show()
        self.usser = User()
        self.usser.show()

    def ROOT(self):
        from User.script_user import root
        global ruut
        ruut = root

    def zapross(self):
        self.ui.tableWidget_2.setRowCount(0)
        self.ui.tableWidget_2.setColumnCount(0)
        if self.ui.comboBox_3.currentIndex() == 0:
            znachenie = str(self.ui.lineEdit_6.text())
            # self.ui.tableWidget_2.setRowCount()
            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/gurnal_start',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['gurnal']:
                if i[2] == znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(com.json()['films'][i[1] - 1][1]))
                    R += 2
        if self.ui.comboBox_3.currentIndex() == 1:
            znachenie = float(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(2)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название", "Дата трансляции"])
            com = requests.post(
                'http://127.0.0.1:5000/gurnal_start',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['gurnal']:
                if i[3] > znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(com.json()['films'][i[1] - 1][1]))
                    self.ui.tableWidget_2.setItem(R, 1, QTableWidgetItem(i[2]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 2:
            znachenie = float(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(2)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название", "Дата трансляции"])
            com = requests.post(
                'http://127.0.0.1:5000/gurnal_start',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['gurnal']:
                if i[3] < znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(com.json()['films'][i[1] - 1][1]))
                    self.ui.tableWidget_2.setItem(R, 1, QTableWidgetItem(i[2]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 3:
            znachenie = float(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(2)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название", "Дата трансляции"])
            com = requests.post(
                'http://127.0.0.1:5000/gurnal_start',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['gurnal']:
                if i[3] == znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(com.json()['films'][i[1] - 1][1]))
                    self.ui.tableWidget_2.setItem(R, 1, QTableWidgetItem(i[2]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 4:
            znachenie = str(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )
            R = 1
            for i in com.json()['films']:
                if i[2].find(znachenie) > -1:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 5:
            znachenie = str(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )
            R = 1
            for i in com.json()['films']:
                if i[3].find(znachenie) > -1:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 6:
            znachenie = float(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['films']:
                if float(i[7]) > znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 7:
            znachenie = float(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['films']:
                if float(i[7]) < znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 8:
            znachenie = float(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['films']:
                if float(i[7]) == znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 9:
            # znachenie = float(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['films']:
                if str(i[9]) == '1':
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 10:
            # znachenie = float(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['films']:
                if str(i[9]) == '0':
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 11:
            # znachenie = float(self.ui.lineEdit_6.text())
            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['films']:
                if str(i[9]) == '2':
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 12:
            znachenie = int(self.ui.lineEdit_6.text())

            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['films']:
                if int(i[5].replace('$', '')) == znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 13:
            znachenie = int(self.ui.lineEdit_6.text())

            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['films']:
                if int(i[5].replace('$', '')) > znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

        if self.ui.comboBox_3.currentIndex() == 14:
            znachenie = int(self.ui.lineEdit_6.text())

            self.ui.tableWidget_2.setColumnCount(1)
            self.ui.tableWidget_2.setHorizontalHeaderLabels(["Название"])
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )

            R = 1
            for i in com.json()['films']:
                if int(i[5].replace('$', '')) < znachenie:
                    self.ui.tableWidget_2.setRowCount(R)
                    R -= 1
                    self.ui.tableWidget_2.setItem(R, 0, QTableWidgetItem(i[1]))
                    R += 2

    def FILMS(self):
        # global cursor, conn
        # conn = sqlite3.connect("bd.db")
        # cursor = conn.cursor()
        try:
            self.ui.comboBox.clear()
            self.ui.comboBox_2.clear()
            self.ui.comboBox_4.clear()
            com = requests.post(
                'http://127.0.0.1:5000/films',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )
            for i in com.json()['films']:
                self.ui.comboBox.addItem(i[1])
                self.ui.comboBox_2.addItem(i[1])
            self.ui.label_2.setText('Название: ' + com.json()['films'][self.ui.comboBox.currentIndex()][1])
            self.ui.label_3.setText('Страна: ' + com.json()['films'][self.ui.comboBox.currentIndex()][2])
            self.ui.label_4.setText('Жанр: ' + com.json()['films'][self.ui.comboBox.currentIndex()][3])
            self.ui.label_5.setText('Длительность: ' + com.json()['films'][self.ui.comboBox.currentIndex()][4])
            self.ui.label_30.setText('Дата выхода: ' + com.json()['films'][self.ui.comboBox.currentIndex()][8])
            self.ui.label_6.setText('Бюджет: ' + com.json()['films'][self.ui.comboBox.currentIndex()][5])
            self.ui.label_7.setText('Рейтинг: ' + str(com.json()['films'][self.ui.comboBox.currentIndex()][7]))
            if com.json()['films'][self.ui.comboBox.currentIndex()][9] == 0:
                self.ui.label_8.setText('Состояние: выбыл из проката')
            elif com.json()['films'][self.ui.comboBox.currentIndex()][9] == 1:
                self.ui.label_8.setText('Состояние: в прокате')
            else:
                self.ui.label_8.setText('Состояние: ожидает проката')
            self.ui.plainTextEdit.setPlainText(com.json()['films'][self.ui.comboBox.currentIndex()][6])

            for i in com.json()['kinogurnal']:
                if i[1] == 1:
                    self.ui.comboBox_4.addItem(i[2])
        except:
            print('Connection error')
            print('1')

    def inf_film(self):
        try:
            com = requests.post(
                'http://127.0.0.1:5000/inf_film',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )
            self.ui.label_2.setText('Название: ' + com.json()['films'][self.ui.comboBox.currentIndex()][1])
            self.ui.label_3.setText('Страна: ' + com.json()['films'][self.ui.comboBox.currentIndex()][2])
            self.ui.label_4.setText('Жанр: ' + com.json()['films'][self.ui.comboBox.currentIndex()][3])
            self.ui.label_5.setText('Длительность: ' + com.json()['films'][self.ui.comboBox.currentIndex()][4])
            self.ui.label_30.setText('Дата выхода: ' + com.json()['films'][self.ui.comboBox.currentIndex()][8])
            self.ui.label_6.setText('Бюджет: ' + com.json()['films'][self.ui.comboBox.currentIndex()][5])
            self.ui.label_7.setText('Рейтинг: ' + str(com.json()['films'][self.ui.comboBox.currentIndex()][7]))
            if com.json()['films'][self.ui.comboBox.currentIndex()][9] == 0:
                self.ui.label_8.setText('Состояние: выбыл из проката')
            elif com.json()['films'][self.ui.comboBox.currentIndex()][9] == 1:
                self.ui.label_8.setText('Состояние: в прокате')
            else:
                self.ui.label_8.setText('Состояние: ожидает проката')
            self.ui.plainTextEdit.setPlainText(com.json()['films'][self.ui.comboBox.currentIndex()][6])
        except:
            print('Connection error')
            print('7')

    def films_mesta_obr_one(self):
        try:
            com = requests.post(
                'http://127.0.0.1:5000/films_mesta_obr_one',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )
            self.ui.comboBox_4.clear()

            for i in com.json()['kinogurnal']:
                if i[1] == self.ui.comboBox_2.currentIndex() + 1:
                    self.ui.comboBox_4.addItem(i[2])

            self.films_mesta_obr()
        except:
            print('Connection error')
            print('8')

    def films_mesta_obr(self):
        try:
            com = requests.post(
                'http://127.0.0.1:5000/films_mesta_obr',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )
            for i in com.json()['mesto']:
                if i[1] == self.ui.comboBox_2.currentIndex() + 1 and i[6] == self.ui.comboBox_4.currentText():
                    KodStroka = eval('self.ui.pushButton_' + str(i[5]) + '.setStyleSheet')
                    if i[4] == 0:
                        KodStroka("background-color: rgb(52, 59, 72);\n"
                                  "border-style: solid;\n"
                                  "border-width: 1px;\n"
                                  "border-radius: 25px;\n"
                                  "border-color: rgb(52, 59, 72);")
                    elif i[4] == 1:
                        KodStroka("background-color: rgb(27, 29, 35);\n"
                                  "border-style: solid;\n"
                                  "border-width: 1px;\n"
                                  "border-radius: 25px;\n"
                                  "border-color: rgb(27, 29, 35);")
        except:
            print('Connection error')
            print('2')

    def obraboka_mest(self, x, y, z):
        self.ROOT()
        if ruut:
            KodStroka = eval('self.ui.pushButton_' + str(z) + '.setStyleSheet')
            inf_mesto = []
            try:
                com = requests.post(
                    'http://127.0.0.1:5000/obraboka_mest_mesto',
                    json={'comboBox': self.ui.comboBox.currentIndex()}
                )

                for i in com.json()['mesto']:
                    if i[1] == self.ui.comboBox_2.currentIndex() + 1 and i[6] == self.ui.comboBox_4.currentText():
                        inf_mesto.append(i)
            except:
                print('Connection error')

            for i in inf_mesto:
                if i[2] == x and i[3] == y:
                    if i[4] == 0:
                        comm = requests.post(
                            'http://127.0.0.1:5000/obraboka_mest_inf_mesto_1',
                            json={'i': i}
                        )
                        KodStroka("background-color: rgb(27, 29, 35);\n"
                                  "border-style: solid;\n"
                                  "border-width: 1px;\n"
                                  "border-radius: 25px;\n"
                                  "border-color: rgb(27, 29, 35);")
                    elif i[4] == 1:
                        commm = requests.post(
                            'http://127.0.0.1:5000/obraboka_mest_inf_mesto_0',
                            json={'i': i}
                        )
                        KodStroka("background-color: rgb(52, 59, 72);\n"
                                  "border-style: solid;\n"
                                  "border-width: 1px;\n"
                                  "border-radius: 25px;\n"
                                  "border-color: rgb(52, 59, 72);")
            commmm = requests.post(
                'http://127.0.0.1:5000/obraboka_mest_inf_mesto_2',
                json={'i': i}
            )
            inf_mesto = []
            for i in commmm.json()['mesto']:
                if i[1] == self.ui.comboBox_2.currentIndex() + 1 and i[6] == self.ui.comboBox_4.currentText():
                    inf_mesto.append(i)
            zanyto = 0
            for i in inf_mesto:
                if i[4] == 1:
                    zanyto += 1
            procent = 90 / 100 * zanyto
            commmmm = requests.post(
                'http://127.0.0.1:5000/obraboka_mest_inf_mesto_3',
                json={'comboBox_2': self.ui.comboBox_2.currentIndex(), 'comboBox_4': self.ui.comboBox_4.currentText(),
                      'procent': procent}
            )
            self.gurnal_start()

    def del_prokat(self):
        self.ROOT()
        if ruut:
            try:
                com = requests.post(
                    'http://127.0.0.1:5000/del_prokat',
                    json={'comboBox': self.ui.comboBox.currentIndex()}
                )
                # print(com.json()['comm'])
                # Получаем данные с таблица Фильмы
                if com.json()['ok']:
                    self.ui.label_8.setText('Состояние: выбыл из проката')
                    self.inf()
                else:
                    msg = QMessageBox()
                    msg.setIcon(QMessageBox.Information)
                    msg.setWindowTitle("Ошибка")
                    msg.setText("Фильм не находится в прокате")
                    okButton = msg.addButton('Хорошо', QMessageBox.AcceptRole)
                    msg.exec()
            except:
                print('Connection error')
                print('3')

    def add_prokat(self):
        self.ROOT()
        if ruut:
            try:
                com = requests.post(
                    'http://127.0.0.1:5000/add_prokat',
                    json={'comboBox': self.ui.comboBox.currentIndex()}
                )
                # print(com.json()['comm'])
                # Получаем данные с таблица Фильмы
                if com.json()['ok']:
                    self.ui.label_8.setText('Состояние: в прокате')
                    self.inf()
                else:
                    msg = QMessageBox()
                    msg.setIcon(QMessageBox.Information)
                    msg.setWindowTitle("Ошибка")
                    msg.setText("Фильм уже находится в прокате")
                    okButton = msg.addButton('Хорошо', QMessageBox.AcceptRole)
                    msg.exec()
            except:
                print('Connection error')
                print('4')

    def gurnal_start(self):
        try:
            com = requests.post(
                'http://127.0.0.1:5000/gurnal_start',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )
            len_gurnal = len(com.json()['gurnal'])
            self.ui.tableWidget.setRowCount(len_gurnal)
            h = 0
            for i in com.json()['gurnal']:
                self.ui.tableWidget.setItem(h, 1, QTableWidgetItem(i[2]))
                self.ui.tableWidget.setItem(h, 2, QTableWidgetItem(str(i[3])))
                h += 1
            h = 0
            for i in com.json()['gurnal']:
                for j in com.json()['films']:
                    if i[1] == j[0]:
                        self.ui.tableWidget.setItem(h, 0, QTableWidgetItem(j[1]))
                h += 1
        except:
            print('Connection error')
            print('5')

    def add_gurnal_zapis(self):
        self.ROOT()
        if ruut:
            self.addzapis = AddZApis()
            self.addzapis.show()

    def seve_gurnal(self):
        self.ROOT()
        data_t = []
        for i in range(int(self.ui.tableWidget.rowCount())):
            data_t.append(str(self.ui.tableWidget.item(i, 1).text()))
        if ruut:
            try:
                try:
                    com = requests.post(
                        'http://127.0.0.1:5000/seve_gurnal',
                        json={'data_t': data_t}
                    )
                except:
                    print('Connection error')
                    print('9')
            except AttributeError:
                self.showNormal()
                self.FILMS()
                self.films_mesta_obr()
                self.gurnal_start()
                self.inf()

    def del_gurnal_zapis(self):
        self.ROOT()
        if ruut:
            try:
                com = requests.post(
                    'http://127.0.0.1:5000/del_gurnal_zapis',
                    json={'lineEdit': self.ui.lineEdit.text()}
                )
                if com.json()['msg']:
                    msg = QMessageBox()
                    msg.setIcon(QMessageBox.Information)
                    msg.setWindowTitle("Ошибка")
                    msg.setText("Данная запись отсуствует")
                    okButton = msg.addButton('Хорошо', QMessageBox.AcceptRole)
                    msg.exec()

                else:
                    self.gurnal_start()
            except:
                print('Connection error')
                print('10')

    def inf(self):
        try:
            com = requests.post(
                'http://127.0.0.1:5000/inf',
                json={'comboBox': self.ui.comboBox.currentIndex()}
            )
            kol_vo_films = len(com.json()['films'])
            ogidanie_prokata = 0
            vibili_prokat = 0
            prokat = 0
            for i in com.json()['films']:
                if i[9] == 0:
                    vibili_prokat += 1
                elif i[9] == 1:
                    prokat += 1
                elif i[9] == 2:
                    ogidanie_prokata += 1

            self.ui.lineEdit_2.setText(str(kol_vo_films))
            self.ui.lineEdit_3.setText(str(ogidanie_prokata))
            self.ui.lineEdit_4.setText(str(vibili_prokat))
            self.ui.lineEdit_5.setText(str(prokat))
        except:
            print('Connection error')
            print('6')

    def printer(self):
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(16)
        font.name = 'Times New Roman'

        name_film = doc.add_paragraph()
        name_film.alignment = 1
        name_film.add_run(films[self.ui.comboBox.currentIndex()][1]).bold = True
        name_film.style = doc.styles['Heading 1']

        obj_styles = doc.styles['Normal']
        obj_font = obj_styles.font
        obj_font.size = Pt(12)
        obj_font.name = 'Times New Roman'

        country = doc.add_paragraph()
        country.add_run('Страна: ').bold = True
        country.add_run(films[self.ui.comboBox.currentIndex()][2])

        Genre = doc.add_paragraph()
        Genre.add_run('Жанр: ').bold = True
        Genre.add_run(films[self.ui.comboBox.currentIndex()][3])

        duration = doc.add_paragraph()
        duration.add_run('Длительность: ').bold = True
        duration.add_run(films[self.ui.comboBox.currentIndex()][4])

        release_date = doc.add_paragraph()
        release_date.add_run('Дата выхода: ').bold = True
        release_date.add_run(films[self.ui.comboBox.currentIndex()][8])

        budget = doc.add_paragraph()
        budget.add_run('Бюджет: ').bold = True
        budget.add_run(films[self.ui.comboBox.currentIndex()][5])

        rating = doc.add_paragraph()
        rating.add_run('Рейтинг: ').bold = True
        rating.add_run(str(films[self.ui.comboBox.currentIndex()][7]))

        if films[self.ui.comboBox.currentIndex()][9] == 0:
            state = doc.add_paragraph()
            state.add_run('Состояние: ').bold = True
            state.add_run('выбыл из проката')
        elif films[self.ui.comboBox.currentIndex()][9] == 1:
            state = doc.add_paragraph()
            state.add_run('Состояние: ').bold = True
            state.add_run('в прокате')
        else:
            state = doc.add_paragraph()
            state.add_run('Состояние: ').bold = True
            state.add_run('ожидает проката')

        description_title = doc.add_paragraph()
        description_title.add_run('Описание: ').bold = True
        description_title.add_run(films[self.ui.comboBox.currentIndex()][6])

        doc.save('inf_film.docx')
        os.startfile('inf_film.docx', 'print')

    def mousePressEvent(self, event):
        self.dragPos = event.globalPos()

    def moveWindow(self, event):
        if UIFunctions.returStatus() == 1:
            UIFunctions.maximize_restore(self)

        if event.buttons() == Qt.LeftButton:
            self.move(self.pos() + event.globalPos() - self.dragPos)
            self.dragPos = event.globalPos()
            event.accept()

    def Button(self):
        btnWidget = self.sender()

        if btnWidget.objectName() == "btn_report":
            self.ui.stackedWidget.setCurrentWidget(self.ui.page_report)
            UIFunctions.resetStyle(self, "btn_report")
            UIFunctions.labelPage(self, "Фильмы")
            btnWidget.setStyleSheet(UIFunctions.selectMenu(btnWidget.styleSheet()))

        if btnWidget.objectName() == "btn_bill":
            self.ui.stackedWidget.setCurrentWidget(self.ui.page_bill)
            UIFunctions.resetStyle(self, "btn_bill")
            UIFunctions.labelPage(self, "Места в зале")
            btnWidget.setStyleSheet(UIFunctions.selectMenu(btnWidget.styleSheet()))

        if btnWidget.objectName() == "btn_2":
            self.ui.stackedWidget.setCurrentWidget(self.ui.page_2)
            UIFunctions.resetStyle(self, "btn_2")
            UIFunctions.labelPage(self, "Киножурнал")
            btnWidget.setStyleSheet(UIFunctions.selectMenu(btnWidget.styleSheet()))

        if btnWidget.objectName() == "btn_3":
            self.ui.stackedWidget.setCurrentWidget(self.ui.page_3)
            UIFunctions.resetStyle(self, "btn_3")
            UIFunctions.labelPage(self, "Запросы")
            btnWidget.setStyleSheet(UIFunctions.selectMenu(btnWidget.styleSheet()))

        if btnWidget.objectName() == "btn":
            self.ui.stackedWidget.setCurrentWidget(self.ui.page)
            UIFunctions.resetStyle(self, "btn")
            UIFunctions.labelPage(self, "Информация")
            btnWidget.setStyleSheet(UIFunctions.selectMenu(btnWidget.styleSheet()))


if __name__ == "__main__":
    app = QApplication(sys.argv)
    QtGui.QFontDatabase.addApplicationFont('fonts/segoeui.ttf')
    QtGui.QFontDatabase.addApplicationFont('fonts/segoeuib.ttf')
    window = MainWindows()
    sys.exit(app.exec_())
